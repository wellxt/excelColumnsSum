import pdfplumber       # pip install pdfplumber
from docx import Document   # 不是使用pip install docx来安装，而是pip install python-docx安装
import os
from multiprocessing import Process     # 多进程，python自带库，无须另行安装

def convertPdf(fileName):
    with pdfplumber.open(fileName) as pdf:
        print("正在处理文件{0}，一共{1}页...".format(fileName,len(pdf.pages)))
        content = ''
        baseName = fileName.split('.')[0]   # 以.号为分隔符，取.号前的名称
        wordName = baseName + '.docx'       # 命名转换后word的全名，包括扩展名
        flag = True

        # 首先用if判断语句删除文件名为wordName的docx文档（避免文件重复）
        if os.path.exists(wordName):
            os.remove(wordName)

        # 把pdf的每一页进行循环
        for i in range(len(pdf.pages)):
            print("正在处理<{0}>第{1}页..".format(baseName,i))
            page = pdf.pages[i]     # 提取第i页的内容
            # 解压第i页的内容
            if page.extract_text() == None:
                print("{0}是图片拼接起来的，所以无法转换".format(fileName))    # 此程序不支持图片拼接起来的PDF，只支持文字版的PDF
                flag = False
                break
            page_content = '\n'.join(page.extract_text().split('\n')[:-1])      # 以换行为分割符
            content = content + page_content
            if os.path.exists(wordName):
                doc = Document(wordName)
            else:
                doc = Document()
            doc.add_paragraph(content)     # 增加段落
            doc.save(wordName)             # 保存WORD文档（处理一页保存一次，以免内存爆表）
            content = ''        # 清空content内容
            print("<{0}>第{1}页处理完成！".format(baseName, i))
        if flag:
            print("{0}处理完成，保存为{1}！！！！".format(fileName, wordName))


if __name__ == '__main__':
    for file in os.listdir('.'):
        # 对当前文件夹中的所有文件进行判断，判断是否是文件，且判断扩展名是否是pdf
        if os.path.isfile(file) and file.split('.')[1] == 'pdf':
            # 使用多进程处理目标函数，args是参数，最后的逗号不能少。
            p = Process(target=convertPdf,args=(file,))
            p.start()
